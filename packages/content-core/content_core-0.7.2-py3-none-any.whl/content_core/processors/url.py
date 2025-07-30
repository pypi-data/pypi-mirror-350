import re
from urllib.parse import urlparse
from io import BytesIO

import aiohttp
import docx
from bs4 import BeautifulSoup, Comment

from content_core.common import ProcessSourceState
from content_core.logging import logger
from content_core.processors.pdf import SUPPORTED_FITZ_TYPES

# future: better extraction methods
# https://github.com/buriy/python-readability
# also try readability: from readability import Document

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

async def _extract_docx_content(docx_bytes: bytes, url: str):
    """
    Extract content from DOCX file bytes.
    """
    try:
        logger.debug(f"Attempting to parse DOCX from URL: {url} with python-docx")
        doc = docx.Document(BytesIO(docx_bytes))
        content_parts = [p.text for p in doc.paragraphs if p.text]
        full_content = "\n\n".join(content_parts)
        
        # Try to get a title from document properties or first heading
        title = doc.core_properties.title
        if not title and doc.paragraphs:
            # Look for a potential title in the first few paragraphs (e.g., if styled as heading)
            for p in doc.paragraphs[:5]: # Check first 5 paragraphs
                if p.style.name.startswith('Heading'):
                    title = p.text
                    break
            if not title: # Fallback to first line if no heading found
                 title = doc.paragraphs[0].text.strip() if doc.paragraphs[0].text.strip() else None

        # If no title found, use filename from URL
        if not title:
            title = urlparse(url).path.split('/')[-1]

        logger.info(f"Successfully extracted content from DOCX: {url}, Title: {title}")
        return {
            "title": title,
            "content": full_content,
            "domain": urlparse(url).netloc,
            "url": url,
        }
    except Exception as e:
        logger.error(f"Failed to process DOCX content from {url}: {e}")
        # Fallback or re-raise, depending on desired error handling
        return {
            "title": f"Error Processing DOCX: {urlparse(url).path.split('/')[-1]}",
            "content": f"Failed to extract content from DOCX: {e}",
            "domain": urlparse(url).netloc,
            "url": url,
        }

async def url_provider(state: ProcessSourceState):
    """
    Identify the provider
    """
    return_dict = {}
    url = state.url
    if url:
        if "youtube.com" in url or "youtu.be" in url:
            return_dict["identified_type"] = "youtube"
        else:
            # remote URL: check content-type to catch PDFs
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.head(url, timeout=10, allow_redirects=True) as resp:
                        mime = resp.headers.get("content-type", "").split(";", 1)[0]
            except Exception as e:
                logger.debug(f"HEAD check failed for {url}: {e}")
                mime = "article"
            if mime in SUPPORTED_FITZ_TYPES:
                return_dict["identified_type"] = mime
            else:
                return_dict["identified_type"] = "article"
    return return_dict


async def extract_url_bs4(url: str):
    """
    Get the title and content of a URL using bs4
    """
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }

        # If URL is actually HTML content
        if url.startswith("<!DOCTYPE html>") or url.startswith("<html"):
            html_content = url
        else:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, headers=headers, timeout=10) as response:
                    response.raise_for_status()
                    # Check content type for DOCX
                    if response.content_type == DOCX_MIME_TYPE:
                        logger.debug(f"Detected DOCX content type for {url}")
                        docx_bytes = await response.read()
                        return await _extract_docx_content(docx_bytes, url)
                    
                    # If not DOCX, proceed as HTML
                    html_content = await response.text()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove unwanted elements
        for element in soup.find_all(
            ["script", "style", "nav", "footer", "iframe", "noscript", "ad"]
        ):
            element.decompose()

        # Remove comments
        for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()

        # Get title
        title = None
        title_tags = [
            soup.find("meta", property="og:title"),
            soup.find("meta", property="twitter:title"),
            soup.find("title"),
            soup.find("h1"),
        ]

        for tag in title_tags:
            if tag:
                if tag.string:
                    title = tag.string
                elif tag.get("content"):
                    title = tag.get("content")
                break

        # Clean up title
        if title:
            title = " ".join(title.split())
            title = re.sub(r"\s*\|.*$", "", title)
            title = re.sub(r"\s*-.*$", "", title)

        # Get content
        content = []

        # Look for main article content
        main_content = None
        content_tags = [
            soup.find("article"),
            soup.find("main"),
            soup.find(class_=re.compile(r"article|post|content|entry|document")),
            soup.find(id=re.compile(r"article|post|content|entry|main")),
        ]

        for tag in content_tags:
            if tag:
                main_content = tag
                break

        if not main_content:
            main_content = soup

        # Process content
        for element in main_content.find_all(
            ["p", "h1", "h2", "h3", "h4", "h5", "h6", "pre", "div"]
        ):
            # Handle code blocks
            if element.name == "pre" or "highlight" in element.get("class", []):
                code_text = element.get_text().strip()
                if code_text:
                    content.append("\n```\n" + code_text + "\n```\n")
                continue

            # Handle regular text
            text = element.get_text().strip()
            if text:
                # Skip if text matches common patterns for navigation/footer
                if re.search(
                    r"copyright|all rights reserved|privacy policy|terms of use",
                    text.lower(),
                ):
                    continue

                content.append(text)

        # Join content with proper spacing
        final_content = "\n\n".join(content)

        # Clean up content
        final_content = re.sub(
            r"\n\s*\n\s*\n", "\n\n", final_content
        )  # Remove extra newlines
        final_content = re.sub(r" +", " ", final_content)  # Normalize whitespace
        final_content = final_content.strip()

        return {
            "title": title,
            "content": final_content,
            "domain": urlparse(url).netloc
            if not url.startswith("<!DOCTYPE html>")
            else None,
            "url": url if not url.startswith("<!DOCTYPE html>") else None,
        }

    except aiohttp.ClientError as e:
        logger.error(f"Failed to fetch URL {url}: {e}")
        return None
    except Exception as e:
        logger.error(f"Failed to process content: {e}")
        return None


async def extract_url_jina(url: str):
    """
    Get the content of a URL using Jina
    """
    async with aiohttp.ClientSession() as session:
        async with session.get(f"https://r.jina.ai/{url}") as response:
            text = await response.text()
            if text.startswith("Title:") and "\n" in text:
                title_end = text.index("\n")
                title = text[6:title_end].strip()
                content = text[title_end + 1 :].strip()
                logger.debug(
                    f"Processed url: {url}, found title: {title}, content: {content[:100]}..."
                )
                return {"title": title, "content": content}
            else:
                logger.debug(
                    f"Processed url: {url}, does not have Title prefix, returning full content: {text[:100]}..."
                )
                return {"content": text}


async def extract_url(state: ProcessSourceState):
    assert state.url, "No URL provided"
    url = state.url
    try:
        result = await extract_url_bs4(url)
        if not result or not result.get("content"):
            logger.debug(
                f"BS4 extraction failed for url {url}, falling back to Jina extractor"
            )
            result = await extract_url_jina(url)
        return result
    except Exception as e:
        logger.error(f"URL extraction failed for URL: {url}")
        logger.exception(e)
        return None
