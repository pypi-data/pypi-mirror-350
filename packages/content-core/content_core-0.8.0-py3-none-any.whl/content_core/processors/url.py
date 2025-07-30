import os
from io import BytesIO
from urllib.parse import urlparse

import aiohttp
import docx
from bs4 import BeautifulSoup
from readability import Document

from content_core.common import ProcessSourceState
from content_core.common.types import warn_if_deprecated_engine
from content_core.logging import logger
from content_core.processors.pdf import SUPPORTED_FITZ_TYPES

DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


async def _extract_docx_content(docx_bytes: bytes, url: str):
    """
    Extract content from DOCX file bytes.
    """
    try:
        logger.debug(f"Attempting to parse DOCX from URL: {url} with python-docx")
        doc = docx.Document(BytesIO(docx_bytes))
        content_parts = [p.text for p in doc.paragraphs if p.text]
        full_content = "\n\n".join(content_parts)

        # Try to get a title from document properties or first heading
        title = doc.core_properties.title
        if not title and doc.paragraphs:
            # Look for a potential title in the first few paragraphs (e.g., if styled as heading)
            for p in doc.paragraphs[:5]:  # Check first 5 paragraphs
                if p.style.name.startswith("Heading"):
                    title = p.text
                    break
            if not title:  # Fallback to first line if no heading found
                title = (
                    doc.paragraphs[0].text.strip()
                    if doc.paragraphs[0].text.strip()
                    else None
                )

        # If no title found, use filename from URL
        if not title:
            title = urlparse(url).path.split("/")[-1]

        logger.info(f"Successfully extracted content from DOCX: {url}, Title: {title}")
        return {
            "title": title,
            "content": full_content,
            "domain": urlparse(url).netloc,
            "url": url,
        }
    except Exception as e:
        logger.error(f"Failed to process DOCX content from {url}: {e}")
        # Fallback or re-raise, depending on desired error handling
        return {
            "title": f"Error Processing DOCX: {urlparse(url).path.split('/')[-1]}",
            "content": f"Failed to extract content from DOCX: {e}",
            "domain": urlparse(url).netloc,
            "url": url,
        }


async def url_provider(state: ProcessSourceState):
    """
    Identify the provider
    """
    return_dict = {}
    url = state.url
    if url:
        if "youtube.com" in url or "youtu.be" in url:
            return_dict["identified_type"] = "youtube"
        else:
            # remote URL: check content-type to catch PDFs
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.head(
                        url, timeout=10, allow_redirects=True
                    ) as resp:
                        mime = resp.headers.get("content-type", "").split(";", 1)[0]
            except Exception as e:
                logger.debug(f"HEAD check failed for {url}: {e}")
                mime = "article"
            if mime in SUPPORTED_FITZ_TYPES:
                return_dict["identified_type"] = mime
            else:
                return_dict["identified_type"] = "article"
    return return_dict


async def extract_url_bs4(url: str) -> dict:
    """
    Get the title and content of a URL using readability with a fallback to BeautifulSoup.

    Args:
        url (str): The URL of the webpage to extract content from.

    Returns:
        dict: A dictionary containing the 'title' and 'content' of the webpage.
    """
    async with aiohttp.ClientSession() as session:
        try:
            # Fetch the webpage content
            async with session.get(url, timeout=10) as response:
                if response.status != 200:
                    raise Exception(f"HTTP error: {response.status}")
                html = await response.text()

            # Try extracting with readability
            try:
                doc = Document(html)
                title = doc.title() or "No title found"
                # Extract content as plain text by parsing the cleaned HTML
                soup = BeautifulSoup(doc.summary(), "lxml")
                content = soup.get_text(separator=" ", strip=True)
                if not content.strip():
                    raise ValueError("No content extracted by readability")
            except Exception as e:
                print(f"Readability failed: {e}")
                # Fallback to BeautifulSoup
                soup = BeautifulSoup(html, "lxml")
                # Extract title
                title_tag = (
                    soup.find("title")
                    or soup.find("h1")
                    or soup.find("meta", property="og:title")
                )
                title = (
                    title_tag.get_text(strip=True) if title_tag else "No title found"
                )
                # Extract content from common content tags
                content_tags = soup.select(
                    'article, .content, .post, main, [role="main"], div[class*="content"], div[class*="article"]'
                )
                content = (
                    " ".join(
                        tag.get_text(separator=" ", strip=True) for tag in content_tags
                    )
                    if content_tags
                    else soup.get_text(separator=" ", strip=True)
                )
                content = content.strip() or "No content found"

            return {
                "title": title,
                "content": content,
            }

        except Exception as e:
            print(f"Error processing URL {url}: {e}")
            return {
                "title": "Error",
                "content": f"Failed to extract content: {str(e)}",
            }


async def extract_url_jina(url: str):
    """
    Get the content of a URL using Jina. Uses Bearer token if JINA_API_KEY is set.
    """
    headers = {}
    api_key = os.environ.get("JINA_API_KEY")
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    async with aiohttp.ClientSession() as session:
        async with session.get(f"https://r.jina.ai/{url}", headers=headers) as response:
            text = await response.text()
            if text.startswith("Title:") and "\n" in text:
                title_end = text.index("\n")
                title = text[6:title_end].strip()
                content = text[title_end + 1 :].strip()
                logger.debug(
                    f"Processed url: {url}, found title: {title}, content: {content[:100]}..."
                )
                return {"title": title, "content": content}
            else:
                logger.debug(
                    f"Processed url: {url}, does not have Title prefix, returning full content: {text[:100]}..."
                )
                return {"content": text}


async def extract_url_firecrawl(url: str):
    """
    Get the content of a URL using Firecrawl.
    Returns {"title": ..., "content": ...} or None on failure.
    """
    try:
        from firecrawl import AsyncFirecrawlApp

        app = AsyncFirecrawlApp(api_key=os.environ.get("FIRECRAWL_API_KEY"))
        scrape_result = await app.scrape_url(url, formats=["markdown", "html"])
        return {
            "title": scrape_result.metadata["title"] or scrape_result.title,
            "content": scrape_result.markdown,
        }

    except Exception as e:
        logger.error(f"Firecrawl extraction error for URL: {url}: {e}")
        return None


async def extract_url(state: ProcessSourceState):
    """
    Extract content from a URL using the engine specified in the state.
    Supported engines: 'auto', 'simple', 'legacy' (deprecated), 'firecrawl', 'jina'.
    """
    assert state.url, "No URL provided"
    url = state.url
    engine = state.engine or "auto"
    warn_if_deprecated_engine(engine)
    try:
        if engine == "auto":
            if os.environ.get("FIRECRAWL_API_KEY"):
                logger.debug(
                    "Engine 'auto' selected: using Firecrawl (FIRECRAWL_API_KEY detected)"
                )
                return await extract_url_firecrawl(url)
            else:
                try:
                    logger.debug("Trying to use Jina to extract URL")
                    return await extract_url_jina(url)
                except Exception as e:
                    logger.error(f"Jina extraction error for URL: {url}: {e}")
                    logger.debug("Falling back to BeautifulSoup")
                    return await extract_url_bs4(url)
        elif engine == "simple" or engine == "legacy":
            # 'legacy' is deprecated alias for 'simple'
            return await extract_url_bs4(url)
        elif engine == "firecrawl":
            return await extract_url_firecrawl(url)
        elif engine == "jina":
            return await extract_url_jina(url)
        elif engine == "docling":
            from content_core.processors.docling import extract_with_docling

            state.url = url
            result_state = await extract_with_docling(state)
            return {"title": None, "content": result_state.content}
        else:
            raise ValueError(f"Unknown engine: {engine}")
    except Exception as e:
        logger.error(f"URL extraction failed for URL: {url}")
        logger.exception(e)
        return None
