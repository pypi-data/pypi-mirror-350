import os
import argparse
import git
import shutil
from pathlib import Path
from docx import Document
import tempfile
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CodebaseToText:
    def __init__(self, input_path, output_path, output_type="txt", verbose=False, exclude_hidden=False, include_exts=None):
        self.input_path = input_path
        self.output_path = output_path
        self.output_type = output_type.lower()
        self.verbose = verbose
        self.exclude_hidden = exclude_hidden
        self.include_exts = [ext.lower() for ext in include_exts] if include_exts else None
        self.temp_folder_path = None
        if self.verbose:
            logger.setLevel(logging.DEBUG)
        else:
            logger.setLevel(logging.INFO)

    def _is_hidden_file(self, file_path):
        path = Path(file_path)
        return any(part.startswith('.') or part.startswith('__') for part in path.parts)

    def _is_included_file(self, file_path):
        if self.include_exts is None:
            return True
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.include_exts

    def _parse_folder(self, folder_path):
        tree = ""
        for root, dirs, files in os.walk(folder_path):
            if self.exclude_hidden:
                dirs[:] = [d for d in dirs if not self._is_hidden_file(os.path.join(root, d))]

            level = root.replace(folder_path, '').count(os.sep)
            indent = ' ' * 4 * (level)
            # Exclude hidden folders from folder structure output
            if self.exclude_hidden and self._is_hidden_file(os.path.basename(root)):
                continue
            tree += '{}{}/\n'.format(indent, os.path.basename(root))
            subindent = ' ' * 4 * (level + 1)
            if self.exclude_hidden:
                files = [f for f in files if not self._is_hidden_file(os.path.join(root, f))]
            if self.include_exts is not None:
                files = [f for f in files if self._is_included_file(f)]
            for f in files:
                tree += '{}{}\n'.format(subindent, f)

        if self.verbose:
            logger.debug(f"The file tree to be processed:\n{tree}")

        return tree

    def _get_file_contents(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            logger.warning(f"Could not decode file {file_path} with utf-8. Skipping content.")
            return ""
        except Exception as e:
            logger.warning(f"Error reading file {file_path}: {e}")
            return ""

    def _process_files(self, path):
        for root, _, files in os.walk(path):
            # Skip .git directories
            if '/.git' in root:
                if self.verbose:
                    logger.debug(f"Ignoring .git directory: {root}")
                continue

            for file in files:
                file_path = os.path.join(root, file)
                # Skip files within .git directory
                if '/.git/' in file_path:
                    if self.verbose:
                        logger.debug(f"Ignoring file in .git directory: {file_path}")
                    continue

                if self.exclude_hidden and self._is_hidden_file(os.path.abspath(file_path)):
                    if self.verbose:
                        logger.debug(f"Ignoring hidden file {file_path}")
                    continue
                if self.include_exts is not None and not self._is_included_file(file):
                    if self.verbose:
                        logger.debug(f"Skipping file due to extension filter: {file_path}")
                    continue
                try:
                    if self.verbose:
                        logger.debug(f"Processing: {file_path}")
                    file_content = self._get_file_contents(file_path)
                    yield f"\n\n{file_path}\n"
                    yield f"File type: {os.path.splitext(file_path)[1]}\n"
                    yield file_content
                    yield f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
                except Exception as e:
                    logger.error(f"Couldn't process {file_path}: {e}")

    def get_text(self):
        folder_structure = ""
        file_content_generator = None
        if self.is_github_repo():
            self._clone_github_repo() # Cleanup is handled within _clone_github_repo now
            folder_structure = self._parse_folder(self.temp_folder_path)
            file_content_generator = self._process_files(self.temp_folder_path)
        else:
            folder_structure = self._parse_folder(self.input_path)
            file_content_generator = self._process_files(self.input_path)

        folder_structure_header = "Folder Structure"
        file_contents_header = "File Contents"
        delimiter = "-" * 50

        # Use a generator to build the text content
        def text_generator():
            yield f"{folder_structure_header}\n{delimiter}\n{folder_structure}\n\n{file_contents_header}\n{delimiter}\n"
            yield from file_content_generator

        return text_generator()

    def get_file(self):
        text_generator = self.get_text()
        if self.output_type == "txt":
            with open(self.output_path, "w", encoding='utf-8') as file:
                for chunk in text_generator:
                    file.write(chunk)
        elif self.output_type == "docx":
            doc = Document()
            # Add folder structure as heading
            doc.add_heading("Folder Structure", level=1)
            # The folder structure is already generated in get_text, but we need it separately for docx
            # Re-parsing here is acceptable as it's fast and doesn't involve file reads
            folder_structure = self._parse_folder(self.temp_folder_path if self.is_github_repo() else self.input_path)
            doc.add_paragraph(folder_structure)
            # Add file contents with file path as heading
            # Iterate through the generator to get file contents
            current_file_path = None
            current_content = []
            for chunk in text_generator:
                if chunk.strip().startswith('/') and '\nFile type:' in chunk:
                    # This is a file path and type header
                    # This is a file path and type header
                    # Add heading for the new file
                    lines = chunk.strip().split('\n')
                    current_file_path = lines[0]
                    doc.add_heading(current_file_path, level=2)
                    current_content = [] # Reset content for the new file
                elif chunk.strip().startswith('---') and chunk.strip().endswith('---'):
                    # This is a file end delimiter, process the accumulated content
                    if current_content: # Only add paragraph if there is content
                         doc.add_paragraph(''.join(current_content).strip())
                    current_file_path = None # Reset file path after processing
                    current_content = [] # Reset content after processing
                else:
                    # This is file content
                    current_content.append(chunk)

            # Add the last file's content if any
            if current_content: # Check if there is any remaining content for the last file
                 doc.add_paragraph(''.join(current_content).strip())

            doc.save(self.output_path)
        else:
            raise ValueError("Invalid output type. Supported types: txt, docx")

    def _clone_github_repo(self):
        try:
            self.temp_folder_path = tempfile.mkdtemp(prefix="github_repo_")
            if self.verbose:
                logger.info(f"Cloning GitHub repository {self.input_path} into {self.temp_folder_path}")
            git.Repo.clone_from(self.input_path, self.temp_folder_path)
            if self.verbose:
                logger.info("GitHub repository cloned successfully.")
        except Exception as e:
            logger.error(f"Error cloning GitHub repository: {e}")
            # Ensure cleanup on failure
            self.clean_up_temp_folder()
            raise e # Re-raise the exception

    def is_github_repo(self):
        return self.input_path.startswith("https://github.com/") or self.input_path.startswith("git@github.com:")

    def clean_up_temp_folder(self):
        if self.temp_folder_path:
            if self.verbose:
                logger.info(f"Cleaning up temporary folder {self.temp_folder_path}")
            shutil.rmtree(self.temp_folder_path)
            self.temp_folder_path = None

class SourceToTextConverter:
    def __init__(self, input_path, output_path, output_type="txt", verbose=False, exclude_hidden=False, include_exts=None):
        self.input_path = input_path
        self.output_path = output_path
        self.output_type = output_type.lower()
        self.verbose = verbose
        self.exclude_hidden = exclude_hidden
        self.include_exts = [ext.lower() for ext in include_exts] if include_exts else None
        self.temp_folder_path = None
        if self.verbose:
            logger.setLevel(logging.DEBUG)
        else:
            logger.setLevel(logging.INFO)

    def _is_hidden_file(self, file_path):
        components = os.path.normpath(file_path).split(os.sep)
        for c in components:
            if c.startswith((".", "__")):
                return True
        return False

    def _is_included_file(self, file_path):
        if self.include_exts is None:
            return True
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.include_exts

    def _parse_folder(self, folder_path):
        tree = ""
        for root, dirs, files in os.walk(folder_path):
            if self.exclude_hidden:
                dirs[:] = [d for d in dirs if not self._is_hidden_file(os.path.join(root, d))]

            level = root.replace(folder_path, '').count(os.sep)
            indent = ' ' * 4 * (level)
            # Exclude hidden folders from folder structure output
            if self.exclude_hidden and self._is_hidden_file(os.path.basename(root)):
                continue
            tree += '{}{}/\n'.format(indent, os.path.basename(root))
            subindent = ' ' * 4 * (level + 1)
            if self.exclude_hidden:
                files = [f for f in files if not self._is_hidden_file(os.path.join(root, f))]
            if self.include_exts is not None:
                files = [f for f in files if self._is_included_file(f)]
            for f in files:
                tree += '{}{}\n'.format(subindent, f)

        if self.verbose:
            logger.debug(f"The file tree to be processed:\n{tree}")

        return tree

    def _get_file_contents(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            logger.warning(f"Could not decode file {file_path} with utf-8. Skipping content.")
            return ""
        except Exception as e:
            logger.warning(f"Error reading file {file_path}: {e}")
            return ""

    def _process_files(self, path):
        content = ""
        for root, _, files in os.walk(path):
            for file in files:
                file_path = os.path.join(root, file)
                if self.exclude_hidden and self._is_hidden_file(os.path.abspath(file_path)):
                    if self.verbose:
                        logger.debug(f"Ignoring hidden file {file_path}")
                    continue
                if self.include_exts is not None and not self._is_included_file(file):
                    if self.verbose:
                        logger.debug(f"Skipping file due to extension filter: {file_path}")
                    continue
                try:
                    if self.verbose:
                        logger.debug(f"Processing: {file_path}")
                    file_content = self._get_file_contents(file_path)
                    content += f"\n\n{file_path}\n"
                    content += f"File type: {os.path.splitext(file_path)[1]}\n"
                    content += f"{file_content}"
                    content += f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
                except Exception as e:
                    logger.error(f"Couldn't process {file_path}: {e}")
        return content

    def get_text(self):
        folder_structure = ""
        file_contents = ""
        if self.is_github_repo():
            self._clone_github_repo()
            folder_structure = self._parse_folder(self.temp_folder_path)
            file_contents = self._process_files(self.temp_folder_path)
            self.clean_up_temp_folder()
        else:
            folder_structure = self._parse_folder(self.input_path)
            file_contents = self._process_files(self.input_path)

        folder_structure_header = "Folder Structure"
        file_contents_header = "File Contents"
        delimiter = "-" * 50

        final_text = f"{folder_structure_header}\n{delimiter}\n{folder_structure}\n\n{file_contents_header}\n{delimiter}\n{file_contents}"

        return final_text

    def get_file(self):
        text = self.get_text()
        if self.output_type == "txt":
            with open(self.output_path, "w", encoding='utf-8') as file:
                file.write(text)
        elif self.output_type == "docx":
            doc = Document()
            # Add folder structure as heading
            doc.add_heading("Folder Structure", level=1)
            folder_structure = self._parse_folder(self.temp_folder_path if self.is_github_repo() else self.input_path)
            doc.add_paragraph(folder_structure)
            # Add file contents with file path as heading
            if self.is_github_repo():
                base_path = self.temp_folder_path
            else:
                base_path = self.input_path
            for root, _, files in os.walk(base_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if self.exclude_hidden and self._is_hidden_file(os.path.abspath(file_path)):
                        continue
                    if self.include_exts is not None and not self._is_included_file(file):
                        continue
                    try:
                        file_content = self._get_file_contents(file_path)
                        relative_path = os.path.relpath(file_path, base_path)
                        doc.add_heading(relative_path, level=2)
                        doc.add_paragraph(file_content)
                    except Exception as e:
                        logger.error(f"Couldn't process {file_path}: {e}")
            doc.save(self.output_path)
        else:
            raise ValueError("Invalid output type. Supported types: txt, docx")

    def _clone_github_repo(self):
        try:
            self.temp_folder_path = tempfile.mkdtemp(prefix="github_repo_")
            if self.verbose:
                logger.info(f"Cloning GitHub repository {self.input_path} into {self.temp_folder_path}")
            git.Repo.clone_from(self.input_path, self.temp_folder_path)
            if self.verbose:
                logger.info("GitHub repository cloned successfully.")
        except Exception as e:
            logger.error(f"Error cloning GitHub repository: {e}")

    def is_github_repo(self):
        return self.input_path.startswith("https://github.com/") or self.input_path.startswith("git@github.com:")

    def clean_up_temp_folder(self):
        if self.temp_folder_path:
            if self.verbose:
                logger.info(f"Cleaning up temporary folder {self.temp_folder_path}")
            shutil.rmtree(self.temp_folder_path)
            self.temp_folder_path = None

def main():
    parser = argparse.ArgumentParser(description="Generate text from codebase.")
    parser.add_argument("--input", help="Input path (folder or GitHub URL)", required=True)
    parser.add_argument("--output", help="Output file path", required=True)
    parser.add_argument("--output-type", help="Output file type (txt or docx)", default="txt")
    parser.add_argument("--exclude-hidden", help="Exclude hidden files and folders", action="store_true")
    parser.add_argument("--include-ext", help="Comma-separated list of file extensions to include (e.g. .py,.js)")
    parser.add_argument("--verbose", help="Show useful information", action="store_true")
    args = parser.parse_args()

    include_exts = None
    if args.include_ext:
        include_exts = [ext.strip().lower() for ext in args.include_ext.split(",")]

    code_to_text = SourceToTextConverter(
        input_path=args.input,
        output_path=args.output,
        output_type=args.output_type,
        verbose=args.verbose,
        exclude_hidden=args.exclude_hidden,
        include_exts=include_exts
    )
    code_to_text.get_file()

if __name__ == "__main__":
    main()
