#!/usr/bin/env python3
#
# Copyright (C) 2025 Lucas Aimaretto / laimaretto@gmail.com, Beatriz Bonafe / bonafencb@gmail.com
# Copyright (C) 2024 Lucas Aimaretto / laimaretto@gmail.com, Beatriz Bonafe / bonafencb@gmail.com , Kathleen Mendonca / kathleencristine20@gmail.com
# Copyright (C) 2023 Lucas Aimaretto / laimaretto@gmail.com
# Copyright (C) 2020 Manuel Saldivar / manuelsaldivar@outlook.com.ar, Lucas Aimaretto / laimaretto@gmail.com
# 
# This is logChecker
# 
# logChecker is free software: you can redistribute it and/or modify
# it under the terms of the 3-clause BSD License.
# 
# logChecker is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY of any kind whatsoever.
# 

import textfsm
import pandas as pd
import glob
import argparse
import sys
import json
import re
from ttp import ttp
import os
import io
import time
import importlib
import traceback

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

LATEST_VERSION = '4.5.11'

DATA_VALUE         = 'Value'
DATA_COMMAND       = '#Command:'
DATA_MAJOR_DWN     = '#majorDown:'
DATA_FLTR_COLS     = '#filterColumns:'
DATA_FLTR_ACTN     = '#filterAction:'

PRE                = 'Pre'
POST               = 'Post'
IDX_PRE_POST       = 'Idx Pre/Post'

INDEX_COL = {
	'sheet' : {
		'position': 0, 'col': 'A:A', 'colName': 'Sheet', 'width': 20,
	},
	'command' : {
		'position': 1, 'col': 'B:B', 'colName': 'Command', 'width': 35,
	},
    'status' : {
		'position': 2, 'col': 'C:C', 'colName' : 'Status', 'width' : 35,
	}
}

RTR_ID = dict(
	name = ['NAME'],
	both = ['NAME','IP'],
	ip   = ['IP']
)

CELL_COLOR = 'black'
CELL_FONT_SIZE = '12'
NO_MATCH = '\n(([N|n]o [M|m]atching [E|e]ntr(y|(ies))( [F|f]ound)?(\.)?\n)|(No.+?: 0\n)|(Number of.+: 0\n))(-+|=+)?$'

D_STATUS = dict(
	no_parsing = dict( #Hay comando, no parsing
		colorTab = '#4092FF', #blue
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### NO Parsing Detected #######',
		shortText = "Can't Parse",
		),
	no_matching_entries = dict(
		colorTab = '#CCCCCC', #gray
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### No Matching Entries #######',
		shortText = 'No Matching Entries',
		),
	no_template = dict( #no hay template, no hay parsing
		colorTab = '#9765FE', #purple
		warnText = '####### NO Template in template folder #######',
		errText  = '####### NO Template in template folder #######',
		shortText = 'No Template',
		),
	no_data = dict( #hay comando, hay parsing, pero sin datos
		colorTab = '#F06AE5', #pink
		warnText = '####### No Data in Command #######',
		errText  = '####### No Data in Command #######',
		shortText = 'No Data in Log',
		),
	ok = dict(
		colorTab = '#37CC73', #green
		warnText = '####### NO POST-TASK CHANGES DETECTED #######',
		errText  = '####### NO MAJOR ERRORS FOUND #######',
		shortText = 'Ok!',
		),
	changes_detected = dict(
		colorTab = '#FFD568', #yellow
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### NO MAJOR ERRORS FOUND #######',
		shortText = 'Warning',
		),
	major_errors = dict(
		colorTab = '#F47F31', #orange
		warnText = '####### CHANGES DETECTED #######',
		errText  = '####### MAJOR ERRORS DETECTED POST-TASK #######',
		shortText = 'Major Errors',
	),
	ambiguity = dict(
		colorTab = '#40FFE8', #teal
		warnText = "####### CAN'T COMPARE. PLEASE USE THE SPECIFIC TEMPLATE #######",
		errText  = None,
		shortText = "Can't compare: use specific template",
	)
)

GENERAL_TEMPL = 'general.template'

GENERAL_TEMPL_LINES = """#Command: .+
#Timos: any
Value Required Lines (.+)

Start
  ^${Lines} -> Record"""

NON_COMMAND_KEYS = ['name','ip','version','hwType','#FINSCRIPT','exit all','','/environment no more']

PLUGIN_ERROR = '''Plugin Error: the plugin must return a dictionary with the structure:
  dictPlugin = {"sheetname_n": {"df": df_n, "valueKeys": valueKeys_n}} for n sheets,
  where df_n must be a dataFrame and valueKeys_n must be a list.
  Or dictPlugin must be None, if there is no need to create new sheets.'''

def readTemplate(fileTemplate, templateFolder, templateEngine):
	'''
	Read the list of templates passed by CSV of textFSM and return template read list (read)

	List of parsed variable names, list of template names

	If fileTemplate is omitted, then all the templates inside the folder are considered.
	'''

	if fileTemplate != '':
		with open(fileTemplate,'r') as f:
			templates = [x.replace('\n','') for x in f.readlines()]
	else:
		if os.path.exists(templateFolder):
			templates = [f.replace(templateFolder,'') for f in glob.glob(templateFolder + '*') if 'majorFile.yml' not in f]
		else: #Si no hay carpeta con templates...
			templateFolder = ''
			templates = []

	d = {}

	d[GENERAL_TEMPL] = {
		'templateColumns':[],
		'commandKey':'',
		'majorDown':['down','dwn'], #En función findMajor, case=False. Aquí no es necesario tener 'Down' y 'Dwn'
		'filterColumns':[],
		'filterAction':None,
		'valueKeys':["Lines"]
	}

	templates.append(GENERAL_TEMPL)

	for i,tmpltName in enumerate(templates):

		d[tmpltName] = {
			'templateColumns':[],
			'commandKey':'',
			'majorDown':['down','dwn'], #En función findMajor, case=False. Aquí no es necesario tener 'Down' y 'Dwn'
			'filterColumns':[],
			'filterAction':None,
			'valueKeys':[]
		}

		if tmpltName == GENERAL_TEMPL:
			tmpltLines = GENERAL_TEMPL_LINES.splitlines()

		else:
			fName = templateFolder+tmpltName
			try:
				with open(fName) as f:
					tmpltLines = f.readlines()
			except:
				print(f'The template file {tmpltName} does not exist inside the folder {templateFolder}.\nPlease check.\nQuitting...')
				quit()

		for line in tmpltLines:

			if templateEngine == 'textFSM':

				h1 = line.find(DATA_VALUE)
				h2 = line.find(DATA_COMMAND)
				h3 = line.find(DATA_MAJOR_DWN)
				h4 = line.find(DATA_FLTR_COLS)
				h5 = line.find(DATA_FLTR_ACTN)
				
				if h1 != -1:
					# We identify here the variables
					col = line.split(' ')[-2]
					d[tmpltName]['templateColumns'].append(col)
					if 'Required' in line or 'Filldown' in line or 'Key' in line:
						d[tmpltName]['valueKeys'].append(col)
				
				if h2 != -1:
					# Here we identify the command
					line = line.replace(DATA_COMMAND + ' ', DATA_COMMAND)
					cmd  = line.split(':')[1].strip('\n')
					d[tmpltName]['commandKey'] = cmd

				if h3 != -1:
					# we add more major words to the list
					line = line.replace(DATA_MAJOR_DWN + ' ', DATA_MAJOR_DWN)
					keys = line.split(':')[1].strip('\n').split(',')
					for key in keys:
						d[tmpltName]['majorDown'].append(key)

				if h4 != -1:
					# We identify the columnes to be filtered
					line = line.replace(DATA_FLTR_COLS + ' ', DATA_FLTR_COLS)
					keys = line.split(':')[1].strip('\n').split(',')
					for key in keys:
						if key not in [None, '', ' ']:
							d[tmpltName]['filterColumns'].append(key)

				if h5 != -1:
					# we identify the action to be performed on the filterd columns
					line = line.replace(DATA_FLTR_ACTN + ' ', DATA_FLTR_ACTN)
					action = line.split(':')[1].strip('\n')
					d[tmpltName]['filterAction'] = action

			if templateEngine == 'ttp':

				h1 = line.find('#Columns: ')
				h2 = line.find('#Command: ')
				h3 = line.find('#majorDown: ')
				h4 = line.find('#filterColumns: ')
				h5 = line.find('#filterAction: ')
				
				if h1 != -1:
					col = line.split(': ')[1].strip('\n').split(",")
					d[tmpltName]['templateColumns'] = col
				
				if h2 != -1:
					cmd = line.split(': ')[1].strip('\n')
					d[tmpltName]['commandKey'] = cmd

				if h3 != -1:
					keys = line.split(': ')[1].strip('\n').split(',')
					for key in keys:
						d[tmpltName]['majorDown'].append(key)

				if h4 != -1:
					# We identify the columnes to be filtered
					keys = line.split(': ')[1].strip('\n').split(',')
					for key in keys:
						if key not in [None, '', ' ']:
							d[tmpltName]['filterColumns'].append(key)

				if h5 != -1:
					# we identify the action to be performed on the filterd columns
					action = line.split(': ')[1].strip('\n')
					d[tmpltName]['filterAction'] = action

		if len(d[tmpltName]['filterColumns']) > 0:

			print(f' The template {tmpltName} has the following columns to be filtered:')
			print(f'  {d[tmpltName]["filterAction"]} the following columns: {", ".join(d[tmpltName]["filterColumns"])}')

			# checking column's names
			x = [col for col in d[tmpltName]['filterColumns'] if col not in d[tmpltName]['templateColumns']]
			if len(x) > 0:
				print(f'There are some columns which are not original variables of the template.')
				print(x)
				print(f'Check the variables names. Quitting...')
				quit()

			# we want to filter columns
			if d[tmpltName]['filterAction'] not in ['include-only','exclude']:
				# we want to filter columns but we have not specified
				# an action to perform
				print(f'The the action to be used has not been properly set.')
				print(f'Please set either "include-only" or "exclude" in the comments section of the template file.\nQuitting...')
				quit()

			if d[tmpltName]['filterAction'] == 'exclude':

				# we check if the filter columns are equal to the templates' columns.
				# if so, chance is we're getting an empty DF. We don't want this.
				if sorted(d[tmpltName]['filterColumns']) == sorted(d[tmpltName]['templateColumns']):
					print(f'The template {tmpltName} has the following columns:')
					print(d[tmpltName]['templateColumns'])
					print(f'Since the action to be performed is "exclude" all the columns will be filtered out and you')
					print(f'will end with an empty table. Quitting...')
					quit()

				# since we are using 'exclude' our final list of filtered columns is the difference
				# between the original 'templateColumns' and the 'filterColumns'
				x = [col for col in d[tmpltName]['templateColumns'] if col not in d[tmpltName]['filterColumns']]
				d[tmpltName]['filterColumns'] = x

			# Only keeps valueKeys columns that are in filterColumns
			d[tmpltName]['valueKeys'] = [col for col in d[tmpltName]['valueKeys'] if col in d[tmpltName]['filterColumns']]

		else:
			# if no filtering columns are defined, we assign those by the original
			# template columns
			d[tmpltName]['filterColumns'] = d[tmpltName]['templateColumns'].copy()

	print(f'##### Successfully Loaded Templates from folder {templateFolder} #####')
	return d

def makeParsed(nomTemplate, routerLog, templateFolder, templateEngine, templateColumns):
	"""
	Parse through textFSM (reading the file again)

	Args:
		nomTemplate (string): name of file containing the textFSM template
		routerLog (string):   logs of router
		templateFolder (string): folder containing the templates
		templateEngine (string): type of templates
		templateColumns (list): columns in the template

	Returns:
		dataframe with parsed results
	"""

	if templateEngine == 'textFSM':

		if nomTemplate == GENERAL_TEMPL:
			template = io.StringIO(GENERAL_TEMPL_LINES) #Para leer correctamente en textfsm.TextFSM(template)
		else:
			template = open(templateFolder + nomTemplate)

		results_template = textfsm.TextFSM(template)
		parsed_results   = results_template.ParseText(routerLog)

		# With list of results, we build a Pandas DataFrame
		parsed_results = pd.DataFrame(parsed_results, columns= templateColumns)

	if templateEngine == 'ttp':

		with open(templateFolder + nomTemplate) as f:
			template = f.read()

		parser = ttp(data=routerLog, template=template)
		parser.parse()

		output = parser.result(format='table')
		parsed_results = output[0][1][0]

		parsed_results = pd.DataFrame(parsed_results, columns= templateColumns)

	return parsed_results

def readLog(logFolder, formatJson):
	"""
	Reads logs, and stores router logs in memory for processing

	Args:
		logFolder (string):  name of folder
		formatJson (string): "yes" or "no"

	Returns: dictionary with logs
	"""

	if formatJson is True:

		ending = '*rx.json'

	else:

		ending = '*rx.txt'

	if sys.platform == "linux" or sys.platform == "linux2" or sys.platform == "darwin":
		# linux

		listContent  = [f for f in glob.glob(logFolder  + ending)]

	elif sys.platform == "win64" or sys.platform == "win32":
		# Windows 64-bit

		listContent  = [f.replace("\\", '/') for f in glob.glob(logFolder  + ending)]
	else:
		print(sys.platform + ": not a valid platform. Quitting....")
		quit()

	d = {}

	if formatJson is True:

		for name in listContent:
			with open(name) as f:
				d[name] = json.load(f)

	else:
	
		for name in listContent:
			with open(name) as f:
				d[name] = f.read()

	print(f'##### Logs Loaded Successfully from folder {logFolder} #####')

	return d

def parseResults(dTmpl, dLog, templateFolder, templateEngine, routerId, useGen, logFolder):
	"""
	Build the Dataframe from textFSM filter, index and router log

	Args:
		dTmpl (dict):           dictionary with info from templates.
		dLog (dict):            dicitonary with logs. Each key is the fileName; the value, is the content of the log.
		templateFolder (str):   folder of templates
		templateEngine:         textFsm or ttp
		routerId:               name, IP or both

	Returns:
		datosEquipo (dict): Dictionary where keys are templateNames. For each key, a DF with parsed results.
	"""
	
	def detParseStatus(datosCmdsLogs,dfTemp):
		"""
		To determine the parseStatus. Options: no_matching_entries, no_parsing, no_data, ok.
		Here, we don't consider the comparision between pre and post logs (statuses: changes_detected and major_errors)
		"""

		parseStatus = 'no_template'
		if len(datosCmdsLogs) > 0 and len(dfTemp) == 0:
			if re.search(NO_MATCH, datosCmdsLogs):
				parseStatus = 'no_matching_entries'
			else:
				parseStatus = 'no_parsing'
		elif len(dfTemp) == 0 and len(datosCmdsLogs) == 0:
			parseStatus = 'no_data'
		else:
			parseStatus = 'ok'
		
		return parseStatus
	
	def writeDfTemp(dfResult, filterCols, orderedColums, routerId, routerName, routerIP, dfTemp):

		# If there are columns to be filtered, we reduced the 
		# size of the DF to that number of columns
		if len(filterCols) > 0:
			dfResult = dfResult[filterCols].copy()

		# We need to define the identification of the router.
		if 'NAME' in RTR_ID[routerId]:
			dfResult['NAME'] = routerName

		if 'IP' in RTR_ID[routerId]:
			dfResult['IP']   = str(routerIP)

		dfResult = dfResult[orderedColums]
		dfTemp = pd.concat([dfTemp, dfResult])

		# It is stored in the dataEquipment dictionary with the key nomTemplate
		# the DF with the data of all routers
		return dfTemp

	def mixAll(dTmpl, datosEquipo, routerId, cmdsLogs, datosCmdsLogs, tmpltName):
		"""
		Function that generates the final DF and obtains parsing status.
		It also updates the datosEquipo dict.
		The counter 'i' is only used when tmpltName is GENERAL_TEMPL
		"""

		if re.match(r"general_\d+", tmpltName):
			templateColumns	= dTmpl[GENERAL_TEMPL]['templateColumns']
			filterCols		= dTmpl[GENERAL_TEMPL]['filterColumns']
			orderedColums	= RTR_ID[routerId] + filterCols
		else:
			templateColumns	= dTmpl[tmpltName]['templateColumns']
			filterCols		= dTmpl[tmpltName]['filterColumns']
			orderedColums	= RTR_ID[routerId] + filterCols

		if tmpltName not in datosEquipo:
			datosEquipo[tmpltName] = {}

		datosEquipo[tmpltName]['command'] = cmdsLogs

		if re.match(r"general_\d+", tmpltName):
			datosEquipo[tmpltName]['template'] = GENERAL_TEMPL
			datosEquipo[tmpltName]['valueKeys'] = ["Lines"]
		else:
			datosEquipo[tmpltName]['template'] = tmpltName
			datosEquipo[tmpltName]['valueKeys'] = dTmpl[tmpltName]['valueKeys']

		if 'dfResultDatos' not in datosEquipo[tmpltName]:
			datosEquipo[tmpltName]['dfResultDatos'] = pd.DataFrame()

		routerLog = cmdsLogs + '\n' + datosCmdsLogs + '\n' #Command and your data

		# We parse results from the key:value association
		# A list is returnd with results
		# to parse, with provide the complete set of columns as defined inside the template: templateColumns
		
		if re.match(r"general_\d+", tmpltName):
			dfResult = makeParsed(GENERAL_TEMPL, routerLog, templateFolder, templateEngine, templateColumns)
		else:
			dfResult = makeParsed(tmpltName, routerLog, templateFolder, templateEngine, templateColumns)

		datosEquipo[tmpltName]['dfResultDatos']	= writeDfTemp(dfResult, filterCols, orderedColums, routerId, routerName, routerIP, datosEquipo[tmpltName]['dfResultDatos'])
		datosEquipo[tmpltName]['parseStatus']	= detParseStatus(datosCmdsLogs, datosEquipo[tmpltName]['dfResultDatos'])

		return datosEquipo		

	datosEquipo		    = {}
	dNoMatchedLog	    = {} #Dictionary similar to dLog, but only with noMatched information
	noMatchedCmdAllRtr	= []

	for idR, routerLogKey in enumerate(dLog.keys()): #To each router
		routerLogFname  = routerLogKey.split("/")[-1]

		routerName		= dLog[routerLogKey]['name']
		routerIP		= dLog[routerLogKey]['ip']

		#The previous versions of taskAutom don't have this information
		try:
			routerVersion	= dLog[routerLogKey]['version']
		except:
			routerVersion	= 'NA'
		try:
			routerHwType	= dLog[routerLogKey]['hwType']
		except:
			routerHwType	= 'NA'

		#For use just keys with command:
		command_keys = [k for k in dLog[routerLogKey].keys() if k not in NON_COMMAND_KEYS]
		# logs is each command that was executed in router, inside json file.
		
		noMatchedCmdPerRtr = command_keys.copy()
		# To control which command have no matching template (here, per router)

		#For each command in command_keys(list)
		for cmdsLogs in command_keys: 
			datosCmdsLogs = dLog[routerLogKey][cmdsLogs] #Logs obtained for each command
			matched_templates = []

			#For each template, we test the match with command, 
			for idT, tmpltName in enumerate(dTmpl.keys()):
				commandKey		= dTmpl[tmpltName]['commandKey']

				# command name in each template file
				prog = re.compile(commandKey)

				# searchKey is the regex match between logs and prog
				match = prog.search(cmdsLogs)

				if match and (tmpltName != GENERAL_TEMPL):
					# If there's a match, we take that command off the list noMatchedCmdPerRtr
					# Important for processing cases that use generic template
					matched_templates.append(tmpltName)
					if cmdsLogs in noMatchedCmdPerRtr:
						noMatchedCmdPerRtr.remove(cmdsLogs)
						
			for tmpltName in matched_templates:
					
				datosEquipo = mixAll(dTmpl, datosEquipo, routerId, cmdsLogs, datosCmdsLogs, tmpltName)

		# Writing dNoMatchedLog for each router. At this point, all attempts to match templates already occurred.
		for cmdsLogs in noMatchedCmdPerRtr:

			if cmdsLogs not in noMatchedCmdAllRtr:
				noMatchedCmdAllRtr.append(cmdsLogs) 
				# Adding to list to no-matched commands, containing information of all routers
			
			#Adding the information to dNoMatchedLog: the dict that only has the cases of no-matched command
			if routerLogKey not in dNoMatchedLog:
				dNoMatchedLog[routerLogKey] = {}

			#This information is required in the following "for"
			dNoMatchedLog[routerLogKey][cmdsLogs] = dLog[routerLogKey][cmdsLogs] 
			dNoMatchedLog[routerLogKey]['ip']	  = dLog[routerLogKey]['ip']    # for writeDfTemp
			dNoMatchedLog[routerLogKey]['name']	  = dLog[routerLogKey]['name']	# for writeDfTemp

	#Processing the no-matched commands
	if useGen == True:
		#If useGen == False, the keys with generic template are not created in datosEquipo dict.

		for idR, routerLogKey in enumerate(dNoMatchedLog.keys()):
			# Basic information of the router, for writeDfTemp
			routerName	= dLog[routerLogKey]['name']
			routerIP	= dLog[routerLogKey]['ip']
			for i,cmdsLogs in enumerate(noMatchedCmdAllRtr):
				# Here, enumerate and noMatchedCmdAllRtr are used to ensure that all no-matched commands,
				# indenpendently of the router, will have the same identification (i).
				# So we can keep the information of the same command together correctly

				# If certain router have datosCmdsLogs, we use this information in this iteration.
				# Otherwise, moves to the other iteration of "for" of commands.
				try:
					datosCmdsLogs = dNoMatchedLog[routerLogKey][cmdsLogs] #Logs obtained for each command
				except:
					continue

				datosEquipo = mixAll(dTmpl, datosEquipo, routerId, cmdsLogs, datosCmdsLogs, f'general_{i}')

	if useGen == False and len(noMatchedCmdAllRtr) > 0:
		print(f'  -ug=no. The following commands will not be parsed: {", ".join(noMatchedCmdAllRtr)}')

	print(f'##### Logs from folder {logFolder} parsed #####')	
	return datosEquipo

def verifyPlugin(pluginFilename):
	"""Verifies the plugin template

	Args:
		pluginFilename (str): Name of config template

	Returns:
		module: The module
	"""

	try:
		if pluginFilename.split(".")[-1] == "py":
			spec = importlib.util.spec_from_file_location("usePlugin",pluginFilename)
			mod  = importlib.util.module_from_spec(spec)
			sys.modules["usePlugin"] = mod
			spec.loader.exec_module(mod)
		else:
			print("Verify extension of the file to be '.py'. Quitting...")
			quit()
	except Exception as e:
		print(e)
		print("----\nError importing plugin. Quitting ...")
		quit()
	
	return mod

def applyPlugin(mod, plg, df_final, dTmplt):
	'''Apply the plugin (from mod) to the data in the df_final. Updates the dTmplt and df_final dictionaries.
	'''	
	try:
		dictPlugin = mod.usePlugin(df_final)

		if dictPlugin == None:
			print(f" {plg} - Does not create a new sheet")
			pass

		elif not (isinstance(dictPlugin,dict)):
			print(PLUGIN_ERROR)
			quit()
		
		else:
			for dpname in dictPlugin:
				if  ('df' not in dictPlugin[dpname] or 'valueKeys' not in dictPlugin[dpname]) or (not (isinstance(dictPlugin[dpname]['df'], pd.DataFrame) and isinstance(dictPlugin[dpname]['valueKeys'], list))):
					print(PLUGIN_ERROR)
					print(f'   Please verify the information in {dpname} dict')
					quit()
	
	except:
		print(f'Plugin Error: Check the structure of your plugin {plg}')
		for e in traceback.format_exc().splitlines()[-3:]:
			print(e)
		quit()

	if dictPlugin is not None:
		for dpname in dictPlugin:
			#Keeping alphanumeric characters and '.', replacing other characters with '_'
			plg_key = plg.split(os.path.sep)[-1]
			plg_key = re.sub(r"[^\w.]", "_", plg_key)
			
			df_final[plg_key+'_'+dpname] = {
				'dfResultDatos' : dictPlugin[dpname]['df'],
				'template'      : plg_key+'_'+dpname,
				'command'       : plg + ' - ' + dpname,
				'valueKeys'     : dictPlugin[dpname]['valueKeys'],
				'parseStatus'   : 'ok',
				'filterColumns' : dictPlugin[dpname]['df'].columns.tolist()
			}

			dTmplt[plg_key+'_'+dpname] = {
				'templateColumns' : dictPlugin[dpname]['df'].columns.tolist(),
				'commandKey'      : '',
				'majorDown'       : ['down','dwn'],
				'filterColumns'   : dictPlugin[dpname]['df'].columns.tolist(),
				'valueKeys'       : dictPlugin[dpname]['valueKeys'],
			}

			print(f' {plg} - {dpname} - New sheet')

	return df_final, dTmplt

def searchDiffAll(datosEquipoPre, datosEquipoPost, dTmplt, routerId, idxComp):
	'''
	Makes a new table, in which it brings the differences between two tables (post-pre)
	'''

	countDif = {}

	def obtain_idx_pre_post(dfCompl,df_pre_post, where):
		'''
		To create new column in dfCompl. Detects the matching index from dfPre or dfPost, for each row.
		'''

		df_pre_post.insert(len(df_pre_post.columns), IDX_PRE_POST, df_pre_post.index) #To get the index in column 'index'
		dfCompl_m = dfCompl[dfCompl['Where'] == where].drop(columns=['Where']) #This column is not in df_pre_post
		df_merged = pd.merge(dfCompl_m,df_pre_post,how='left',on=dfCompl_m.columns.tolist())
		
		if where == 'left_only':
			df_merged.insert(len(df_merged.columns)-1,'Where',PRE)
		else:
			df_merged.insert(len(df_merged.columns)-1,'Where',POST)

		return df_merged

	for tmpltName in datosEquipoPre.keys():
		if tmpltName not in countDif:
			countDif[tmpltName] = {}

		# In datosEquipoPre[tmpltName], tmpltName can be generic_0 and that key does not exist in the templates dictionary (dTmplt).
		# That is why it is necessary to store the template in datosEquipoPre[tmpltName]['template'], so that it can then be related
		# to dTmplt[template] (in the examples of the generic template, template == 'generic.template').

		template	= datosEquipoPre[tmpltName]['template']
		filterCols	= dTmplt[template]['filterColumns']
		dfPre       = datosEquipoPre[tmpltName]['dfResultDatos'].reset_index(drop=True)
		dfPost      = datosEquipoPost[tmpltName]['dfResultDatos'].reset_index(drop=True)

		if template != GENERAL_TEMPL:

			dfUnion = pd.merge(dfPre, dfPost, how='outer', indicator='Where').drop_duplicates()
			dfInter = dfUnion[dfUnion.Where=='both']
			dfCompl = dfUnion[~(dfUnion.isin(dfInter))].dropna(axis=0, how='all').drop_duplicates()

			if len(dfCompl) > 0 and idxComp == True: #Add index from pre-post table

				dfComplPre = obtain_idx_pre_post(dfCompl,dfPre,'left_only')
				dfComplPost = obtain_idx_pre_post(dfCompl,dfPost,'right_only')
				dfCompl = pd.concat([dfComplPre,dfComplPost],ignore_index=True)
			else:
				dfCompl['Where'] = dfCompl['Where'].str.replace('left_only',PRE)
				dfCompl['Where'] = dfCompl['Where'].str.replace('right_only',POST)

		elif template == GENERAL_TEMPL and dfPre.shape == dfPost.shape:

			rtrId   = RTR_ID[routerId][0] # This is so, because the key to identify the router can either be its name or IP; check what to do when routerId == 'both'

			dfCompl = pd.DataFrame(columns=dfPre.columns) # We build an empty DF. Works better when routerId == 'both'

			for rtrName in dfPre[rtrId].unique():
				
				tempPre   = dfPre[dfPre[rtrId] == rtrName]
				tempPost  = dfPost[dfPost[rtrId] == rtrName]
				tempComp = tempPre.compare(tempPost, result_names=(PRE,POST), align_axis=0, keep_equal=True).reset_index() #Reset index to remove multiindex
				tempComp = tempComp.rename(columns={'level_1':'Where','level_0':IDX_PRE_POST})
				tempComp[rtrId] = rtrName
				dfCompl = pd.concat([dfCompl,tempComp])
			
			#Maintaining the standard with the results from specific template
			dfCompl = dfCompl.reindex(columns=[col for col in dfCompl.columns if col != IDX_PRE_POST] + [IDX_PRE_POST])

		# When using general template and the dfs from pre and post are != in size, the comparision doesn't work 
		# very well with the options above.

		else:
			datosEquipoPost[tmpltName]['parseStatus'] = 'ambiguity'
			dfCompl = pd.DataFrame(columns=dfPre.columns)

		orderedColums = RTR_ID[routerId] + filterCols

		countDif[tmpltName]['dfResultDatos'] = dfCompl.sort_values(by = orderedColums)
		countDif[tmpltName]['valueKeys'] = datosEquipoPre[tmpltName]['valueKeys']

	print(f'##### Comparisions between logs completed #####')

	return countDif

def findMajor(count_dif, dTmplt, routerId, datosEquipoPre):
	'''
	Makes a table from the results of searching for Major errors in the post table define in yml file for specific template,\n
	or down if is not define the words for the template, which are not in the Pre table
	'''

	countDown  = {}
	
	for tmpltName in count_dif.keys():
		if tmpltName not in countDown:
			countDown[tmpltName] = {}

		df = pd.DataFrame()
		template = datosEquipoPre[tmpltName]['template']

		for majorWord in dTmplt[template]['majorDown']:

			filterCols = dTmplt[template]['filterColumns']

			if 'Where' in count_dif[tmpltName]['dfResultDatos'].columns:

				df1 = count_dif[tmpltName]['dfResultDatos'][count_dif[tmpltName]['dfResultDatos']['Where']==POST]
				
				if len(df1) > 0:
					df1 = df1[df1.apply(lambda r: r.str.contains(majorWord, case=False).any(), axis=1)]
				else:
					df1 = pd.DataFrame(columns=count_dif[tmpltName]['dfResultDatos'].columns)

				df = pd.concat([df, df1])

				df = df.sort_values(by = RTR_ID[routerId] + filterCols)

		df = df.reset_index(drop=True)
		countDown[tmpltName]['dfResultDatos'] = df

	return countDown

def makeTable(datosEquipoPre, datosEquipoPost):
	'''
	Sort the table pre and post to present in Excel
	'''

	df_all          = {}
	datosEquipoPre1 = datosEquipoPre.copy()
	
	for tmpltName in datosEquipoPre.keys():
		if tmpltName not in df_all:
				df_all[tmpltName] = {}

		datosEquipoPre1[tmpltName]['##']='##'

		dfPre1 = datosEquipoPre1[tmpltName]['dfResultDatos'].reset_index(drop=True)
		dfPost = datosEquipoPost[tmpltName]['dfResultDatos'].reset_index(drop=True)
		df_all[tmpltName]['dfResultDatos']	= pd.concat([dfPre1, dfPost], axis=1, keys=('Pre-Check', 'Post-Check'))

		df_all[tmpltName]['parseStatus']	= datosEquipoPost[tmpltName]['parseStatus']
		df_all[tmpltName]['command']		= datosEquipoPre[tmpltName]['command']

	return df_all

def constructExcel(df_final, count_dif, searchMajor, folderLog):
	"""
	Sort the data and format creating the Excel
	_summary_

	Args:
		df_final (pandas): DataFrame with pre and post data
		count_dif (pandas): DataFrame with only differences
		searchMajor (pandas): DataFrame with only errors
		folderLog (string): name of the folder
	"""

	def diff_colors(valueKeys,dfDiff,start_row,worksheet,end_col):

		dfDiff["idx_temp"] = dfDiff[valueKeys].agg('-'.join, axis=1) #Creation column with keys
		id_router = dfDiff[dfDiff.columns[0]].unique() #First column value (name or IP)

		for id in id_router: #Selecting data by router
			df_f_name = dfDiff[dfDiff[dfDiff.columns[0]] == id]

			templ_keys = df_f_name["idx_temp"].unique() 

			for key in templ_keys: #Filtering pre-post pairs for each key
				df_f_keys = df_f_name[df_f_name["idx_temp"] == key]
				row_pair = df_f_keys.index.tolist() #Indices of the pre-post pair rows

				first_pair = dfDiff.iloc[row_pair[0]] #First line of pre-post pair

				if len(row_pair) > 1: # If there's a pre AND post pair
					second_pair = dfDiff.iloc[row_pair[1]]
				else:
					second_pair = None #Case where the row only appears in pre or post

				if second_pair is not None:
					for col in dfDiff.columns:
						if (col not in ['Where', IDX_PRE_POST]) and first_pair[col] != second_pair[col]:
							col_idx = dfDiff.columns.get_loc(col)

							worksheet.conditional_format( #Coloring
								start_row + row_pair[0], col_idx +1, start_row + row_pair[1], col_idx +1,
								{'type': 'no_blanks', 'format': workbook.add_format({'bg_color': '#FFC7CE'}) }
							)
							worksheet.conditional_format( #Black line below, until the end of dfDiff
								start_row + row_pair[1], 2, start_row + row_pair[1], end_col,
								{'type': 'no_blanks', 'format': workbook.add_format({'bottom': 1}) }
							)
							worksheet.conditional_format( #Black line below, until the end of dfDiff for blanks cells
								start_row + row_pair[1], 2, start_row + row_pair[1], end_col,
								{'type': 'blanks', 'format': workbook.add_format({'bottom': 1}) }
							)
							worksheet.conditional_format( #Fixing the first column
								start_row + row_pair[1], 1, start_row + row_pair[1], 1,
								{'type': 'no_blanks', 'format': workbook.add_format({'bottom': 1,'left':1}) }
							)

				else: #Just one pre/post line
					worksheet.conditional_format( #Fixing the first column
						start_row + row_pair[0], 1, start_row + row_pair[0], 1,
						{'type': 'no_blanks', 'format': workbook.add_format({'bottom': 1,'left':1}) }
					)
					worksheet.conditional_format( #Black line below, until the end of dfDiff
						start_row + row_pair[0], 2, start_row + row_pair[0], end_col,
						{'type': 'no_blanks', 'format': workbook.add_format({'bottom': 1,'top': 1}) }
					)
					worksheet.conditional_format( #Black line below, until the end of dfDiff for blanks cells
						start_row + row_pair[0], 2, start_row + row_pair[0], end_col,
						{'type': 'blanks', 'format': workbook.add_format({'bottom': 1,'top': 1}) }
					)

		return

	fileName  = folderLog[:-1] + ".xlsx"

	writer    = pd.ExcelWriter(fileName, engine='xlsxwriter') #creates instance of an excel workbook
	workbook  = writer.book

	# Create index tab
	indexSheet = workbook.add_worksheet('index')

	print('\nSaving Excel')

	for idx,template in enumerate(df_final.keys()):

		dfData  = df_final[template]['dfResultDatos']
		dfDiff  = count_dif[template]['dfResultDatos'].reset_index(drop=True)
		dfMajor = searchMajor[template]['dfResultDatos']
		dfParseStatus = df_final[template]['parseStatus']

		sheet_name = template.replace('nokia_sros_','')
		sheet_name = sheet_name.replace('.template','')
		sheet_name = sheet_name.replace('_template','')
		sheet_name = sheet_name.replace('.ttp','')
		sheet_name = sheet_name.replace('.','_')

		if len(sheet_name) > 31:
			sheet_name = sheet_name[:31]

		# Selecting Tab's color and error messages
		if dfParseStatus not in ['ok','changes_detected','major_errors']:
			output = dfParseStatus
		elif len(dfMajor) == 0 and len(dfDiff) == 0:
			output = 'ok'
		elif len(dfMajor) == 0 and len(dfDiff) != 0:
			output = 'changes_detected'
		elif len(dfMajor) != 0:
			output = 'major_errors'

		# cell format
		cell_format  = workbook.add_format({'color': CELL_COLOR, 'font_size': CELL_FONT_SIZE, 'fg_color': D_STATUS[output]['colorTab'], 'align': 'center', 'border': 1 ,'bold': True})

		srcCol   = 'A'+str(idx+1)

		# Building index
		for k, i_dict in INDEX_COL.items():
			indexSheet.write(0,i_dict['position'],i_dict['colName'], workbook.add_format({'font_size':CELL_FONT_SIZE,'align':'center','border':1,'bold':True}))
			indexSheet.set_column(i_dict['col'],i_dict['width'])

		indexSheet.write_url(idx+1,0, 'internal:'+sheet_name+'!A1', string=sheet_name)
		indexSheet.write(idx+1,1, df_final[template]['command'])
		indexSheet.write(idx+1,2, D_STATUS[output]['shortText'], cell_format)

		# Creating Tab
		worksheet = workbook.add_worksheet(sheet_name)
		worksheet.set_tab_color(D_STATUS[output]['colorTab'])
		writer.sheets[sheet_name] = worksheet
		dfData.to_excel(writer, sheet_name=sheet_name, startrow=0, startcol=0) #Creates Excel File
		worksheet.write_url('A1', 'internal:index!A1', string='Index')

		# Changes Section
		if len(dfDiff) > 0 or output=='ambiguity':
			srcCol   = 'A'+str(len(dfData)+5)
			dstCol   = 'J'+str(len(dfData)+5)
			colRange = srcCol + ':' + dstCol
			warnTex  = D_STATUS[output]['warnText']
			worksheet.merge_range(colRange, warnTex, cell_format)

			if len(dfDiff) > 0:
				valueKeys = count_dif[template]['valueKeys']
				dfDiff.to_excel(writer, sheet_name=sheet_name, startrow=len(dfData)+6, startcol=0)
				start_row = len(dfData) + 7
				end_col = len(dfDiff.columns)

				if len(valueKeys) > 0:
					diff_colors(valueKeys,dfDiff,start_row,worksheet,end_col)

		# Major Error Section
		if len(dfMajor) > 0:
			srcCol   = 'A'+str((len(dfData)+(len(dfDiff)))+9)
			dstCol   = 'J'+str((len(dfData)+(len(dfDiff)))+9)
			colRange = srcCol + ':' + dstCol
			errText   = warnTex  = D_STATUS[output]['errText']
			worksheet.merge_range(colRange, errText, cell_format)
			if len(dfMajor) > 0:
				dfMajor.to_excel(writer, sheet_name=sheet_name, startrow=(len(dfData)+(len(dfDiff)))+10, startcol=0)

		print('#',idx,template)
	
	writer.close() #saves workbook to file in python file directory

def renderAtp(dictParam):
	"""[Generates a ATP based on the json logs obtained from taskAutom.]

	Args:
		dictParam

	Returns:
		None
	"""

	preFolder  = dictParam['preFolder']
	postFolder = dictParam['postFolder']

	jsonFilesPre = [preFolder+x for x in os.listdir(preFolder) if '.json' in x and x != '00_report.json']
	if postFolder != '':
		jsonFilesPos = [postFolder+x for x in os.listdir(postFolder) if '.json' in x and x != '00_report.json']

	job0docx = "./ATP.docx"

	print("\nGenerating ATP: " + job0docx)

	myDoc    = docx.Document()
	myStyles = myDoc.styles

	styleConsole = myStyles.add_style('Console', WD_STYLE_TYPE.PARAGRAPH)
	styleConsole.font.name = 'Courier'
	styleConsole.font.size = Pt(9)
	styleConsole.paragraph_format.keep_together = True

	styleConsole.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	#styleConsole.paragraph_format.line_spacing = Pt(10)
	#styleConsole.paragraph_format.line_spacing = .2
	styleConsole.paragraph_format.space_after = Pt(2)

	myDoc.add_heading('ATP', 0)

	charset_allowed = [chr(c) for c in range(32,127)] + ['\n']

	if preFolder != '':

		docMainTitle = myDoc.add_paragraph('Pre-Check')
		docMainTitle.style = myDoc.styles['Heading 1']
		docMainTitle.paragraph_format.line_spacing = 1.5

		for f in jsonFilesPre:
			
			with open(f) as myFile:
				
				logs  = json.load(myFile)
				keys  = [x for x in logs.keys() if 'show' in x]

				routerTitle = f'Router {logs["name"]} ({logs["ip"]})'

				docRouterTitle = myDoc.add_paragraph(routerTitle)
				docRouterTitle.style = myDoc.styles['Heading 2']
				docRouterTitle.paragraph_format.line_spacing = 1.5

				for key in keys:
					showTitle   = key.rstrip('\n').lstrip('\n')
					showContent = ''.join([x for x in logs[key] if x in charset_allowed]).rstrip('\n').lstrip('\n')

					docShowTitle = myDoc.add_paragraph(showTitle)
					docShowTitle.style = myDoc.styles['Heading 3']
					docShowTitle.paragraph_format.line_spacing = 1.5

					docShowContent = myDoc.add_paragraph(showContent)
					docShowContent.style = myDoc.styles['Console']

	if postFolder != '':

		docMainTitle = myDoc.add_paragraph('Post-Check')
		docMainTitle.style = myDoc.styles['Heading 1']
		docMainTitle.paragraph_format.line_spacing = 1.5

		for f in jsonFilesPos:
			
			with open(f) as myFile:
				
				logs  = json.load(myFile)
				keys  = [x for x in logs.keys() if 'show' in x]

				routerTitle = f'Router {logs["name"]} ({logs["ip"]})'

				docRouterTitle = myDoc.add_paragraph(routerTitle)
				docRouterTitle.style = myDoc.styles['Heading 2']
				docRouterTitle.paragraph_format.line_spacing = 1.5

				for key in keys:
					showTitle   = key.rstrip('\n').lstrip('\n')
					showContent = ''.join([x for x in logs[key] if x in charset_allowed]).rstrip('\n').lstrip('\n')

					docShowTitle = myDoc.add_paragraph(showTitle)
					docShowTitle.style = myDoc.styles['Heading 3']
					docShowTitle.paragraph_format.line_spacing = 1.5

					docShowContent = myDoc.add_paragraph(showContent)
					docShowContent.style = myDoc.styles['Console']

	myDoc.save(job0docx)

	print("ATP done...")

def fncRun(dictParam):

	preFolder          = dictParam['preFolder']
	postFolder         = dictParam['postFolder']
	csvTemplate        = dictParam['csvTemplate']
	formatJson         = dictParam['formatJson']
	templateFolder     = dictParam['templateFolder']
	templateEngine     = dictParam['templateEngine']
	templateFolderPost = dictParam['templateFolderPost']
	routerId           = dictParam['routerId']
	genAtp             = dictParam['genAtp']
	idxComp            = dictParam['idxComp']
	useGen             = dictParam['useGen']
	usePlugin          = dictParam['usePlugin']

	if sys.platform == "win64" or sys.platform == "win32":
		templateFolder = templateFolder.replace('/', '\\')
		if templateFolderPost != '':
			templateFolderPost = templateFolderPost.replace('/','\\')

	if preFolder != '' and postFolder == '':

		dTmplt = readTemplate(csvTemplate, templateFolder, templateEngine)
		dLog   = readLog(preFolder, formatJson)

		df_final    = parseResults(dTmplt, dLog, templateFolder, templateEngine, routerId, useGen, preFolder)
		if len(usePlugin)>0:
			print("##### Plugins: #####")
			for plg in usePlugin:
				mod = verifyPlugin(plg)
				df_final, dTmplt = applyPlugin(mod, plg, df_final, dTmplt)

		count_dif   = {}
		searchMajor = {}

		for tmpltName in df_final.keys():
			if tmpltName not in count_dif:
				count_dif[tmpltName] = {}
			count_dif[tmpltName]['dfResultDatos']   = pd.DataFrame(columns=df_final[tmpltName]['dfResultDatos'].columns)
			if tmpltName not in searchMajor:
				searchMajor[tmpltName] = {}
			searchMajor[tmpltName]['dfResultDatos'] = pd.DataFrame(columns=df_final[tmpltName]['dfResultDatos'].columns)

		constructExcel(df_final, count_dif, searchMajor, preFolder)

		if genAtp is True:
			renderAtp(dictParam)

	elif preFolder != '' and postFolder != '':

		if templateFolder == templateFolderPost:
			dTmpltPre  = readTemplate(csvTemplate, templateFolder, templateEngine)
			dTmpltPost = readTemplate(csvTemplate, templateFolderPost, templateEngine)
		elif templateFolder != '' and templateFolderPost == '':
			templateFolderPost = templateFolder
			dTmpltPre  = readTemplate(csvTemplate, templateFolder, templateEngine)
			dTmpltPost = readTemplate(csvTemplate, templateFolderPost, templateEngine)
		else:
			dTmpltPre  = readTemplate(csvTemplate, templateFolder, templateEngine)
			dTmpltPost = readTemplate(csvTemplate, templateFolderPost, templateEngine)

		dLogPre  = readLog(preFolder, formatJson)
		dLogPost = readLog(postFolder, formatJson)
			
		datosEquipoPre  = parseResults(dTmpltPre,  dLogPre,  templateFolder,     templateEngine, routerId, useGen, preFolder)
		datosEquipoPost = parseResults(dTmpltPost, dLogPost, templateFolderPost, templateEngine, routerId, useGen, postFolder)
		
		if len(usePlugin)>0:
			print("##### Plugins: #####")
			for plg in usePlugin:
				mod = verifyPlugin(plg)
				datosEquipoPre, dTmpltPre   = applyPlugin(mod, plg, datosEquipoPre, dTmpltPre)
				datosEquipoPost, dTmpltPost = applyPlugin(mod, plg, datosEquipoPost, dTmpltPost)

		count_dif       = searchDiffAll(datosEquipoPre, datosEquipoPost, dTmpltPre, routerId, idxComp)

		searchMajor     = findMajor(count_dif, dTmpltPre, routerId, datosEquipoPre)
		df_final        = makeTable(datosEquipoPre, datosEquipoPost)

		constructExcel(df_final, count_dif, searchMajor, postFolder)

		if genAtp is True:
			renderAtp(dictParam)

	elif preFolder == '':
		print('No PRE folder defined. Please Verify.')



def main():
	start_time = time.time()

	parser1 = argparse.ArgumentParser(description='Log Analysis', prog='PROG', usage='%(prog)s [options]')
	parser1.add_argument('-pre', '--preFolder',     type=str, required=True, help='Folder with PRE Logs. Must end in "/"',)
	parser1.add_argument('-post','--postFolder' ,   type=str, default='',    help='Folder with POST Logs. Must end in "/"',)
	parser1.add_argument('-csv', '--csvTemplate',   type=str, default='',    help='CSV with list of templates names to be used in parsing. If the file is omitted, then all the templates inside --templateFolder, will be considered for parsing. Default=None.')
	parser1.add_argument('-json', '--formatJson',   type=str, default='yes', choices=['yes','no'], help='logs in json format: yes or no. Default=yes.')
	parser1.add_argument('-tf', '--templateFolder', type=str, default='Templates/', help='Folder where templates reside. Used both for PRE and POST logs. Default=Templates/')
	parser1.add_argument('-tf-post', '--templateFolderPost', type=str, default='',  help='If set, use this folder of templates for POST logs.')
	parser1.add_argument('-te', '--templateEngine', type=str, default='textFSM', choices=['ttp','textFSM'], help='Engine for parsing. Default=textFSM.')
	parser1.add_argument('-ri', '--routerId',       type=str, default='name',choices=['name','ip','both'],  help='Router Id to be used within the tables in the Excel report. Default=name.')
	parser1.add_argument('-ga', '--genAtp',         type=str, default='no',  choices=['no','yes'], help='Generate ATP document in docx format, based on the contents of the json files from taskAutom. Default=no')
	parser1.add_argument('-ic','--idxComp',         type=str, default='no',  choices=['yes','no'], help='Adds new column (Idx Pre/Post) in changes detected table with . Default=no')
	parser1.add_argument('-ug','--useGen',          type=str, default='yes', choices=['yes','no'], help='Using generic template. If -ug=no, logChecker only use the templates indicated in the -tf and -tf-post folder. Default=yes')
	parser1.add_argument('-up','--usePlugin',type=str, default='',help="Additional plugins for manipulation of parsed information, creating new sheets. One plugin, use -up plugin1.py . For indicate a folder containing all the plugins: -up plugins/ . Default='' ")
	parser1.add_argument('-v' ,'--version',         help='Version', action='version', version=f'(c) 2025 - Version: {LASTEST_VERSION}' )

	args = parser1.parse_args()

	if args.usePlugin == '':
		usePlugin = ''
	elif args.usePlugin.endswith(('/', '\\')):
		usePlugin = glob.glob(args.usePlugin + '*.py')
	else:
		usePlugin = args.usePlugin.split(',')

	dictParam = dict(
		preFolder          = args.preFolder,
		postFolder         = args.postFolder,
		csvTemplate        = args.csvTemplate,
		formatJson         = True if args.formatJson == 'yes' else False,
		templateFolder     = args.templateFolder,
		templateEngine     = args.templateEngine,
		templateFolderPost = args.templateFolderPost,
		routerId           = args.routerId,
		genAtp             = True if args.genAtp == 'yes' else False,
		idxComp            = True if args.idxComp == 'yes' else False,
		useGen             = True if args.useGen == 'yes' else False,
		usePlugin          = usePlugin,
	)

	fncRun(dictParam)

	print(f'\nTotal running time: {time.time()-start_time:.2f} seconds')

### To be run from the python shell
if __name__ == '__main__':
	main()