from docx import Document

def validate_name(name_string):
    if not isinstance(name_string, str) or not name_string.strip():
        return False
    for char in name_string:
        if not (char.isalpha() or char.isspace()):
            return False
    return True

try:
    document = Document('TestCase.docx')
    if not document.tables:
        table = document.add_table(rows=1, cols=3)
    else:
        table = document.tables[0]
except Exception as e:
    print(f"Предупреждение: не удалось открыть 'TestCase.docx' ({e}). Создается новый документ.")
    document = Document()
    table = document.add_table(rows=1, cols=3)

document = Document('TestCase.docx')
table = document.tables[0]
row_cells = table.add_row().cells
row_cells[0].text = 'Проверка формата ФИО'
row_cells[1].text = 'Имя должно содержать только буквы и пробелы.'
row_cells[2].text = 'Прошла успешно!' if validate_name("Иванов") else 'Ошибка!'
document.save('TestCase.docx')