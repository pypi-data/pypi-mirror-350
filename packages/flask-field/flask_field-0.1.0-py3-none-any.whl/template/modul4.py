import tkinter as tk
from tkinter import messagebox
import requests
from docx import Document

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        
        self.title("Приложение для проверки данных")
        self.geometry("400x200")
        
        self.label = tk.Label(self, text="Клиентские данные:")
        self.label.pack()
        
        self.get_data_button = tk.Button(self, text="Получить данные", command=self.fetch_data)
        self.get_data_button.pack(pady=10)
        
        self.test_result_button = tk.Button(self, text="Отправить результат теста", command=self.send_test_result)
        self.test_result_button.pack(pady=10)
    
    def fetch_data(self):

        try:
            response = requests.get("http://localhost:4444/TransferSimulator/fullName")
            if response.status_code == 200:
                data = response.json()["value"] 
                self.label.config(text=f"Клиентские данные: {data}")  
            else:
                messagebox.showerror("Ошибка", f"Не удалось получить данные ({response.status_code})")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Возникла проблема при получении данных: {e}")
    
    def send_test_result(self):

        client_full_name = self.label.cget("text").split(":")[1].strip()  
        
        allowed_chars = set('абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ- ')
        
        invalid_chars = [ch for ch in client_full_name if ch not in allowed_chars]
        
        result_valid = True if not invalid_chars else False
        invalid_message = ", ".join(invalid_chars) if invalid_chars else ""
        
        self.save_results(client_full_name, result_valid, invalid_message)
        
        messagebox.showinfo("Результат проверки",
                            f"Имя: {client_full_name}\n"
                            f"Статус: {'Корректно' if result_valid else 'Некорректно'}\n"
                            f"Недопустимые символы: {invalid_message}")
    
    def save_results(self, client_full_name, result_valid, invalid_message):
        try:
            try:
                doc = Document("TestCase.docx")
            except Exception:
                doc = Document()
            
            if not doc.tables:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ФИО Клиента'
                hdr_cells[1].text = 'Валидность'
                hdr_cells[2].text = 'Недопустимые символы'
            
            row_cells = doc.add_paragraph().add_run()
            table = doc.tables[-1]
            row_cells = table.add_row().cells
            row_cells[0].text = client_full_name
            row_cells[1].text = 'Да' if result_valid else 'Нет'
            row_cells[2].text = invalid_message
            
            doc.save("TestCase.docx")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить результаты: {e}")

if __name__ == "__main__":
    app = App()
    app.mainloop()