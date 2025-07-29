from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 새 Word 문서 생성
doc = Document()

# "Normal" 스타일을 왼쪽 정렬로 설정
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# 문단 정렬을 수동으로 왼쪽으로 지정 (스타일 적용은 수동으로 문단마다 해야 함)
def add_left_aligned_paragraph(text):
    p = doc.add_paragraph(text, style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

# 제목
doc.add_heading('LaTeX 수식 왼쪽 정렬용 reference.docx', level=0)
add_left_aligned_paragraph('이 문서는 Pandoc 변환 시 LaTeX 수식이 왼쪽 정렬되도록 설계된 참조 문서입니다.')

# 인라인 수식 예시
doc.add_heading('1. 인라인 수식 예시', level=1)
add_left_aligned_paragraph('예: 뉴턴의 제2법칙은 F = ma 로 표현됩니다.')

# 블록 수식 예시
doc.add_heading('2. 블록 수식 예시', level=1)
add_left_aligned_paragraph('E = mc²')
add_left_aligned_paragraph('x = (-b ± √(b² - 4ac)) / 2a')

# 정렬된 다중 수식 예시
doc.add_heading('3. 정렬된 수식 예시', level=1)
add_left_aligned_paragraph('y = ax + b')
add_left_aligned_paragraph('z = cx + d')

# 수식 번호 매기기 예시
doc.add_heading('4. 수식 번호 매기기 예시', level=1)
table = doc.add_table(rows=1, cols=2)
table.cell(0, 0).text = 'x = (-b ± √(b² - 4ac)) / 2a'
table.cell(0, 1).text = '(1)'
for cell in table.row_cells(0):
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 수식 입력 방법 안내
doc.add_heading('5. Word에서 수식 입력 방법', level=1)
add_left_aligned_paragraph('1. 상단 메뉴에서 [삽입] 클릭')
add_left_aligned_paragraph('2. [수식] 선택')
add_left_aligned_paragraph('3. LaTeX 문법 또는 수식 입력 도구 사용')

# 저장
# 저장
output_path = "reference_3.docx"
doc.save(output_path)

output_path
