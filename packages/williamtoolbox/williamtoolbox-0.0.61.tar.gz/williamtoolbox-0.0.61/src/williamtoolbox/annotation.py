import re
import os
from typing import List, Dict, Optional, Union
import byzerllm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx import Document
from pydantic import BaseModel
import json
from openai import AsyncOpenAI
from loguru import logger
from byzerllm.utils.client import code_utils
from williamtoolbox.storage.json_file import load_config, load_models_from_json, load_rags_from_json
from autocoder.rag.relevant_utils import FilterDoc


class Annotation(BaseModel):
    text: str
    comment: str


class DocText(BaseModel):
    doc_name: str
    doc_text: str
    annotations: List[Annotation]


class DocPath(BaseModel):
    doc_path: str


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])


def extract_annotations_from_docx(file_path: str) -> List[Dict[str, str]]:
    # Define namespace manually
    NAMESPACE = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    document = Document(file_path)
    annotations = []

    # Extract comments from the document
    comments = {}
    if document.part.rels:
        for rel in document.part.rels.values():
            if rel.reltype == COMMENTS_REL_TYPE:
                comments_part = rel.target_part
                comments_xml = parse_xml(comments_part.blob)
                for comment in comments_xml.findall(".//w:comment", namespaces=NAMESPACE):
                    comment_id = comment.attrib.get(f'{{{NAMESPACE["w"]}}}id')
                    if comment_id:
                        comment_text = ''.join(comment.itertext()).strip()
                        comments[comment_id] = comment_text

    # Match comments to paragraphs
    for paragraph in document.paragraphs:
        paragraph_xml = paragraph._element
        for comment_start in paragraph_xml.findall(".//w:commentRangeStart", namespaces=NAMESPACE):
            comment_id = comment_start.attrib.get(f'{{{NAMESPACE["w"]}}}id')
            if comment_id and comment_id in comments:
                annotations.append({
                    'text': paragraph.text.strip(),
                    'comment': comments[comment_id]
                })

    return annotations


def extract_annotations(text: str) -> List[Dict[str, str]]:
    '''
    Extract annotations from the text.
    Args:
        text: The text with annotations in the format [[[text]]] and <<<comment>>>
    Returns:
        A list of dictionaries with keys 'text' and 'comment'
    '''
    annotations = []
    pattern = r'\[\[\[(.*?)\]\]\]\s*<<<(.*?)>>>'
    matches = re.finditer(pattern, text)

    for match in matches:
        annotations.append({
            'text': match.group(1).strip(),
            'comment': match.group(2).strip()
        })

    return annotations


def process_docx_files(directory: str) -> List[DocText]:
    """
    遍历指定目录中的所有 .docx 文件，提取文本和注释

    Args:
        directory: 要遍历的目录路径

    Returns:
        包含所有文档文本和注释的 DocText 对象列表
    """
    doc_texts = []

    for root, _, files in os.walk(directory):
        for file in files:
            # 跳过隐藏文件
            if file.startswith(('.', '~$')):
                continue

            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                try:
                    text = extract_text_from_docx(file_path)
                    annotations = extract_annotations_from_docx(file_path)
                    doc_texts.append(DocText(
                        doc_name=file_path,
                        doc_text=text,
                        annotations=[Annotation(**a) for a in annotations]
                    ))
                except Exception as e:
                    print(f"Error processing file {file_path}: {str(e)}")
                    continue

    return doc_texts


@byzerllm.prompt()
def generate_annotations(text: str, examples: List[DocText]) -> str:
    '''
    给定一个文本，我们会对里面的特定内容做批注。

    下面是历史批注内容：

    <history>
    {% for example in examples %}
    原始文本：
    <text>
    {{ example.doc_text }}
    </text>

    批注：
    <annotations>
    {% for annotation in example.annotations %} 
    [[[{{ annotation.text }}]]]<<<{{ annotation.comment }}>>>
    {% endfor %}
    </annotations>
    {% endfor %}
    </history>


    请参考历史批注，模仿批注的：
    1. 语气和风格
    2. 批注的角度
    3. 关注的内容点      

    下面是等待批注的文本：
    <text>
    {{ text }}
    </text>

    请你理解文本的含义，模仿上面示例生成对带批注的内容进行批注。
    规则：
    1. 用 [[[]]] 括住需要批注的文本,紧跟着用 <<<>>> 括住对该文本的批注内容，注意，两者中间不要有任何其他字符。
    2. 确保 [[[]]] 括住的需要批注的文本必须完全匹配原文本，不能有任何字符差异。    
    '''


@byzerllm.prompt()
def query_rag(text: str):
    '''    
    下面是等待批注的文档：
    <text>
    {{ text }}
    </text>    

    判断上面的文档是否和等待批注的文档是同一类内容。
    '''


async def chat_with_model(model_name: str, messages: List[Dict[str, str]]) -> str:
    """与指定模型进行聊天"""
    try:
        # 加载配置
        config = await load_config()
        models = await load_models_from_json()
        
        if model_name not in models:
            raise ValueError(f"Model {model_name} not found")
            
        model_info = models[model_name]
        if model_info["status"] != "running":
            raise ValueError(f"Model {model_name} is not running")
            
        # 获取OpenAI服务配置
        openai_server = config.get("openaiServerList", [{}])[0]
        host = openai_server.get("host", "localhost")
        port = openai_server.get("port", 8000)
        if host == "0.0.0.0":
            host = "127.0.0.1"

        base_url = f"http://{host}:{port}/v1"
        client = AsyncOpenAI(base_url=base_url, api_key="xxxx")

        # 调用模型
        response = await client.chat.completions.create(
            model=model_name,
            messages=messages,
            stream=False,
            max_tokens=4*1024
        )
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error in chat_with_model: {str(e)}")
        raise

async def chat_with_rag(rag_name: str, messages: List[Dict[str, str]]) -> str:
    """与指定RAG进行聊天"""
    try:
        rags = await load_rags_from_json()
        
        if rag_name not in rags:
            raise ValueError(f"RAG {rag_name} not found")
            
        rag_info = rags[rag_name]
        if rag_info["status"] != "running":
            raise ValueError(f"RAG {rag_name} is not running")
            
        host = rag_info.get("host", "localhost")
        port = rag_info.get("port", 8000)
        if host == "0.0.0.0":
            host = "127.0.0.1"

        base_url = f"http://{host}:{port}/v1"
        client = AsyncOpenAI(base_url=base_url, api_key="xxxx")

        # 调用RAG
        response = await client.chat.completions.create(
            model=rag_info.get("model", "deepseek_chat"),
            messages=messages,
            stream=False,
            max_tokens=4*1024
        )
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error in chat_with_rag: {str(e)}")
        raise

async def auto_generate_annotations(rag_name: str, doc: str, model_name: str = "default_model") -> DocText:
    # 使用 chat_with_rag 替换 query_rag
    logger.info(f"开始处理文档，文档长度: {len(doc)}")
    
    # 1. 调用 RAG 获取相似文档
    prompt = query_rag.prompt(text=doc[:10000])
    logger.info(f"RAG 查询 prompt:\n{prompt}")

    final_query = json.dumps({
        "query": prompt,
        "only_contexts": True
    },ensure_ascii=False)
    
    rag_response = await chat_with_rag(rag_name, [
        {"role": "user", "content": final_query}
    ])
    logger.info(f"RAG 返回结果:\n{rag_response}")
    docs: List[FilterDoc] = [FilterDoc(**doc) for doc in json.loads(rag_response)]    
    
    # 2. 加载示例文档
    examples = []
    for temp_doc in docs:
        file_path = temp_doc.source_code.module_name  
        try:                
            # clean ##File: /Users/allwefantasy/Downloads/论文/陈梦初稿-已批注.docx.json#chunk1  
            if file_path.startswith("##File: "):
                file_path = file_path[len("##File: "):]
            if "#chunk" in file_path:
                file_path = file_path[:file_path.index("#chunk")]

            with open(file_path, 'r', encoding='utf-8') as file:
                v = json.loads(file.read())
            doc_text = DocText(doc_name=file_path, doc_text=v["doc_text"], annotations=[Annotation(**a) for a in v["annotations"]])
            examples.append(doc_text)
            logger.info(f'成功加载示例文档 {file_path}, 包含 {len(doc_text.annotations)} 条注释')
                
        except Exception as e:
            logger.error(f'加载文档 {file_path} 失败: {str(e)}')            
            continue
    
    # 3. 生成注释
    annotation_prompt = generate_annotations.prompt(text=doc, examples=examples[0:1])
    logger.info(f"生成注释 prompt:\n{annotation_prompt}")
    
    model_response = await chat_with_model(model_name, [
        {"role": "user", "content": annotation_prompt}
    ])
    logger.info(f"模型生成注释响应:\n{model_response}")
    
    # 4. 提取注释
    annotations = extract_annotations(model_response)
    logger.info(f"提取到 {len(annotations)} 条注释")
    for idx, anno in enumerate(annotations, 1):
        logger.info(f"注释 {idx}:")
        logger.info(f"  文本: {anno['text']}")
        logger.info(f"  批注: {anno['comment']}")
    
    return DocText(doc_name="", doc_text=doc, annotations=annotations)



def add_annotations_to_docx(file_path: str, annotations: List[Annotation]) -> None:
    """
    Add annotations to a Word document using Spire.Doc.
    
    Args:
        file_path: Path to the Word document
        annotations: List of Annotation objects containing text and comment pairs
    """
    from spire.doc import Document
    from spire.doc import Comment
    from spire.doc import Paragraph
    from spire.doc import TextSelection
    import re
    # 创建文档对象并加载文件
    doc = Document()
    doc.LoadFromFile(file_path)
    
    # 遍历所有批注
    for annotation in annotations:
        # 查找包含目标文本的段落
        text_to_find = annotation.text
        comment_text = annotation.comment
        
        # 遍历文档中的所有段落
        for i in range(doc.Sections.Count):
            section = doc.Sections.get_Item(i)
            for j in range(section.Paragraphs.Count):
                paragraph = section.Paragraphs.get_Item(j)
                # 查找目标文本在段落中的位置
                text = paragraph.Text
                if text_to_find in text:
                    # 创建批注对象
                    comment = paragraph.AppendComment(comment_text)
                    comment.Format.Author = "Auto Annotation"
                    break
    
    # 保存修改后的文档
    output_path = file_path.replace(".docx", "_annotated.docx")
    doc.SaveToFile(output_path)
    doc.Close()


def add_annotations_to_docx_v2(file_path: str, annotations: List[Annotation]) -> None:
    """
    Add annotations to a Word document using Aspose.Words.
    
    Args:
        file_path: Path to the Word document
        annotations: List of Annotation objects containing text and comment pairs
    """
    import aspose.words as aw
    from datetime import datetime
    
    # 创建文档对象并加载文件
    doc = aw.Document(file_path)
    builder = aw.DocumentBuilder(doc)
    
    # 遍历所有批注
    for annotation in annotations:
        # 查找包含目标文本的段落
        text_to_find = annotation.text
        comment_text = annotation.comment
        
        # 查找目标文本在文档中的位置
        options = aw.replacing.FindReplaceOptions()
        options.direction = aw.replacing.FindReplaceDirection.FORWARD
        options.match_case = True
        
        # 查找并添加批注
        nodes = doc.get_child_nodes(aw.NodeType.RUN, True)
        for node in nodes:
            if text_to_find in node.get_text():
                # 创建批注对象
                comment = aw.Comment(doc, "Auto Annotation", "AA", datetime.now())
                comment.paragraphs.add(aw.Paragraph(doc))
                comment.first_paragraph.runs.add(aw.Run(doc, comment_text))
                
                # 获取当前段落
                current_paragraph = node.get_ancestor(aw.NodeType.PARAGRAPH)
                if current_paragraph is not None:
                    # 将批注添加到段落中
                    comment_range_start = aw.CommentRangeStart(doc, comment.id)
                    comment_range_end = aw.CommentRangeEnd(doc, comment.id)
                    
                    # 插入批注范围
                    current_paragraph.insert_before(comment_range_start, node)
                    current_paragraph.insert_after(comment_range_end, node)
                    
                    # 将批注添加到文档
                    current_paragraph.append_child(comment)
                break
    
    # 保存修改后的文档
    output_path = file_path.replace(".docx", "_annotated_v2.docx")
    doc.save(output_path)
